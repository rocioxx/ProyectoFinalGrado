from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

GOLD  = RGBColor(0xD4, 0xAF, 0x37)
RED   = RGBColor(0xB2, 0x22, 0x22)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY  = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

def h1(text):
    p = doc.add_heading(level=1)
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = GOLD
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)

def h2(text):
    p = doc.add_heading(level=2)
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RED
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)

def body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.bold  = bold
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(4)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)

def tabla(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'B22222')
        tcPr.append(shd)
    for r_i, row_data in enumerate(rows):
        row = t.rows[r_i + 1]
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
        if r_i % 2 == 0:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FDF6E3')
                tcPr.append(shd)
    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])
    doc.add_paragraph()

# ── PORTADA ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MANUAL DE JUGADOR')
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('carDungeon')
run.font.size = Pt(38)
run.font.bold = True
run.font.color.rgb = RED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nJuego de cartas · Versión 1.0\n')
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Desarrollado por Rocío Parra Roldán\n2025-2026')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_page_break()

# ── 1. QUÉ ES ─────────────────────────────────────────────────────────────────
h1('1. ¿Qué es carDungeon?')
body('carDungeon es un juego de cartas en el que eres un aventurero atrapado en '
     'una mazmorra oscura. Para avanzar debes tomar decisiones deslizando cartas '
     'hacia la izquierda o la derecha. Cada elección tiene consecuencias: puedes '
     'ganar vida, perder tiempo, encontrar enemigos o descubrir tesoros.')
body('Tu objetivo es llegar al fondo de la mazmorra y derrotar al Nigromante. '
     'Pero cuidado: si cualquiera de tus cuatro estadísticas llega a cero, '
     'la partida termina.')

# ── 2. INICIO DE SESIÓN ───────────────────────────────────────────────────────
h1('2. Cómo iniciar sesión')
body('Al abrir la aplicación verás la pantalla de portada. Pulsa el botón '
     '"Entrar" para acceder al login.')
body('Tienes dos formas de identificarte:')
bullet('Con tu correo electrónico y contraseña.')
bullet('Con tu cuenta de Google (botón "Continuar con Google").')
body('Si es la primera vez que usas la app con ese correo, el juego te pedirá '
     'que elijas un nombre de jugador y una contraseña para crear tu cuenta.')

# ── 3. PANTALLA DE JUEGO ──────────────────────────────────────────────────────
h1('3. La pantalla de juego')
body('Una vez dentro, verás la pantalla principal con los siguientes elementos:')

tabla(
    ['Elemento', 'Dónde está', 'Para qué sirve'],
    [
        ['Carta',              'Centro de la pantalla',         'Describe la situación actual. Deslízala para decidir.'],
        ['Panel de stats',     'Esquina superior',              'Muestra tus cuatro estadísticas con barras de color.'],
        ['Panel de misiones',  'Esquina superior izquierda',    'Lista las cinco misiones secundarias de la partida.'],
        ['Barra de vida enemigo', 'Parte superior, en combate', 'Aparece solo cuando luchas contra un enemigo.'],
        ['Botón ruleta',       'Parte inferior',                'Salta la carta actual sin consecuencias (gasta tiempo).'],
        ['Texto de opciones',  'Bordes izquierdo y derecho',   'Te indica qué pasará si deslizas en cada dirección.'],
    ],
    col_widths=[3.5, 4, 7.5]
)

# ── 4. LAS ESTADÍSTICAS ───────────────────────────────────────────────────────
h1('4. Las estadísticas')
body('Tu personaje tiene cuatro estadísticas. Si cualquiera llega a 0, pierdes la partida.')

tabla(
    ['Estadística', 'Valor inicial', 'Máximo', 'Qué representa'],
    [
        ['❤  Vida',    '50', '50',  'Tu salud. Bajas cuando recibes daño en combate o eventos negativos.'],
        ['⚔  Poder',  '10', '10',  'Tu fuerza de ataque. El daño que haces es exactamente tu Poder actual.'],
        ['🍀  Suerte', '50', '100', 'Influye en las probabilidades de los eventos de azar (cofres, etc.).'],
        ['⏳  Tiempo', '100','100', 'El tiempo que te queda antes de que la mazmorra te atrape.'],
    ],
    col_widths=[3, 2.5, 2.5, 7]
)
body('Consejo: el Poder es especialmente importante. Con poco Poder harás muy poco '
     'daño en combate y tardarás mucho en matar a los enemigos, perdiendo más Vida '
     'y Tiempo en el proceso.', italic=True)

# ── 5. CÓMO JUGAR ─────────────────────────────────────────────────────────────
h1('5. Cómo jugar')

h2('Deslizar cartas')
body('Cada carta presenta una situación. Lee el texto y decide:')
bullet('Desliza a la IZQUIERDA para elegir la opción de la izquierda.')
bullet('Desliza a la DERECHA para elegir la opción de la derecha.')
body('En los bordes de la pantalla puedes ver un texto que te adelanta qué '
     'consecuencia tendrá cada dirección antes de decidir.')

h2('La ruleta')
body('Si no te gusta la carta actual y es saltable, puedes pulsar el botón de '
     'ruleta en la parte inferior. Esto descarta la carta y te da una nueva, '
     'pero consume Tiempo. Úsala con cuidado.')

h2('El cofre')
body('Cuando encuentras un cofre puedes abrirlo (derecha) o ignorarlo (izquierda). '
     'Si lo abres, hay dos posibles resultados:')
bullet('Poción de suerte: +10 Suerte. ¡Buena noticia!')
bullet('Trampa: -8 Poder. Ten cuidado.')
body('La probabilidad de conseguir la poción depende de tu Suerte actual.')

# ── 6. ENEMIGOS ───────────────────────────────────────────────────────────────
h1('6. Los enemigos')
body('En la mazmorra encontrarás cinco tipos de enemigos. Cuando te enfrentes a '
     'uno, verás su barra de vida en la parte superior de la pantalla. '
     'Desliza las cartas para atacar.')

tabla(
    ['Enemigo', 'Vida', 'Dificultad', 'Consejo'],
    [
        ['Slime',            '20',  'Fácil',    'El más sencillo. Buen calentamiento.'],
        ['Esqueleto',        '30',  'Fácil',    'Sin complicaciones si tienes buen Poder.'],
        ['Esqueleto Armado', '40',  'Media',    'Aguanta más. Mantén el Poder alto.'],
        ['Troll',            '60',  'Difícil',  'Puede hacerte mucho daño. Asegúrate de tener Vida suficiente.'],
        ['Nigromante',       '100', 'Jefe final','A veces golpea el doble. Es el objetivo final del juego.'],
    ],
    col_widths=[3.5, 1.5, 2.5, 7.5]
)
body('El daño que haces por ataque es exactamente tu valor de Poder actual. '
     'Si tienes Poder 2, solo haces 2 de daño por golpe.')

# ── 7. MISIONES ───────────────────────────────────────────────────────────────
h1('7. Misiones secundarias')
body('Tienes cinco misiones que puedes completar durante la partida. '
     'No son obligatorias para ganar, pero se muestran en el resumen final.')

tabla(
    ['Misión', 'Cómo completarla'],
    [
        ['Derrotar al Nigromante',     'Llega al final de la mazmorra y gana el combate final.'],
        ['Llegar a la puerta final',   'Encuentra y atraviesa la puerta sellada al fondo de la mazmorra.'],
        ['Vencer todos los combates',  'Derrota al menos 3 enemigos en la misma partida.'],
        ['Encontrar la cantimplora',   'Encuentra el evento de la cantimplora y recógela.'],
        ['Hablar con el aventurero',   'Encuentra al aventurero misterioso y habla con él.'],
    ],
    col_widths=[5, 10]
)

# ── 8. GANAR Y PERDER ─────────────────────────────────────────────────────────
h1('8. Ganar y perder')

h2('Cómo ganar')
body('Derrota al Nigromante al final de la mazmorra. Verás la pantalla de '
     'victoria con un resumen de las misiones completadas.')

h2('Cómo perder')
body('La partida termina en derrota si cualquiera de tus estadísticas llega a 0:')
bullet('Vida = 0: has muerto en combate o por un evento fatal.')
bullet('Poder = 0: te has quedado sin fuerza para seguir.')
bullet('Suerte = 0: la mala suerte te ha consumido.')
bullet('Tiempo = 0: la mazmorra te ha atrapado antes de llegar al final.')
body('Al perder verás la pantalla de derrota con el resumen de misiones completadas '
     'y los botones para volver al menú o salir del juego.')

# ── 9. CONSEJOS ───────────────────────────────────────────────────────────────
h1('9. Consejos')
bullet('Lee siempre los textos de los bordes antes de deslizar: te adelantan las consecuencias.')
bullet('Conserva el Poder. Sin él, los combates se eternizan y gastas demasiada Vida y Tiempo.')
bullet('Usa la ruleta solo cuando la carta sea realmente mala. Gasta Tiempo.')
bullet('En combate, ataca siempre que puedas. Cuanto antes acabes, menos daño recibirás.')
bullet('El Nigromante a veces golpea el doble. Llega con Vida alta al combate final.')
bullet('La Suerte mejora tus probabilidades en los cofres y otros eventos de azar.')
bullet('Si encuentras la cantimplora o al aventurero, interactúa con ellos: dan ventajas.')

# ── GUARDAR ───────────────────────────────────────────────────────────────────
output = r'C:\ProyectoFinalGrado\Manual_Jugador_carDungeon.docx'
doc.save(output)
print(f'Manual guardado en: {output}')
